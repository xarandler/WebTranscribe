import streamlit as st
import whisper
import sounddevice as sd
import numpy as np
import wave
import threading
import queue
import time
import tempfile
from docx import Document
from io import BytesIO

def seconds_to_srt_time(seconds: float) -> str:
    """Convert seconds to SRT time format HH:MM:SS,mmm."""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    millis = int((seconds - int(seconds)) * 1000)
    return f"{hours:02}:{minutes:02}:{secs:02},{millis:03}"

def generate_word_file(text: str) -> BytesIO:
    """Generate a DOCX file from the given text and return a BytesIO buffer."""
    doc = Document()
    doc.add_heading("Transcription", level=1)
    doc.add_paragraph(text)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def record_audio(filename: str, stop_event: threading.Event, samplerate: int = 44100):
    """
    Record audio from the default microphone until stop_event is set.
    Save the recorded audio as a 16-bit mono WAV file.
    """
    q = queue.Queue()

    def callback(indata, frames, time_info, status):
        if status:
            print(f"Sounddevice status: {status}")
        q.put(indata.copy())

    frames = []
    with sd.InputStream(samplerate=samplerate, channels=1, callback=callback):
        while not stop_event.is_set():
            while not q.empty():
                frames.append(q.get())

    if frames:
        audio_data = np.concatenate(frames, axis=0)
    else:
        audio_data = np.empty((0, 1))
    with wave.open(filename, 'wb') as wf:
        wf.setnchannels(1)   # Mono
        wf.setsampwidth(2)   # 16-bit PCM
        wf.setframerate(samplerate)
        wf.writeframes((audio_data * 32767).astype(np.int16).tobytes())

@st.cache_resource  # For Streamlit 1.22+; if using an older version, use @st.cache(allow_output_mutation=True)
def load_whisper_model(model_size: str):
    """Load and cache the Whisper model."""
    return whisper.load_model(model_size)

def start_recording(audio_filename: str):
    """Callback to start recording audio."""
    st.session_state.stop_event.clear()
    threading.Thread(
        target=record_audio,
        args=(audio_filename, st.session_state.stop_event),
        daemon=True
    ).start()
    st.session_state.recording = True

def stop_recording():
    """Callback to stop recording audio."""
    st.session_state.stop_event.set()
    st.session_state.recording = False
    time.sleep(1)  # Allow time for the WAV file to finalize

def main():
    st.set_page_config(
        page_title="OpenAI Whisper Audio Transcription",
        page_icon="🎙️",
        layout="wide"
    )
    st.title("🎙️ OpenAI Whisper Audio Transcription")

    # ---------------------
    # Sidebar settings
    # ---------------------
    model_size = st.sidebar.selectbox(
        "Model Size",
        ["tiny", "base", "small", "medium", "large"],
        index=2
    )
    language_mode = st.sidebar.radio(
        "Language Mode",
        ("Auto-detect", "Force Swedish"),
        index=0
    )
    translation_mode = st.sidebar.radio(
        "Translation Mode",
        ("Transcribe only", "Translate to English", "Translate to Swedish"),
        index=0
    )

    # Load Whisper model (cached).
    model = load_whisper_model(model_size)
    task = "transcribe" if translation_mode == "Transcribe only" else "translate"

    # ---------------------
    # Prepare session state for recording
    # ---------------------
    audio_filename = "recorded_audio.wav"
    if "recording" not in st.session_state:
        st.session_state.recording = False
    if "stop_event" not in st.session_state:
        st.session_state.stop_event = threading.Event()

    # ---------------------
    # Recording Section
    # ---------------------
    st.subheader("Record Audio")
    if not st.session_state.recording:
        st.button("🎙️ Start Recording",
                  on_click=start_recording,
                  args=(audio_filename,))
        st.info("Click the button above to start recording.")
    else:
        st.warning("🔴 Recording in progress...")
        st.button("🛑 Stop Recording", on_click=stop_recording)
    
    recorded_file_path = None
    if (not st.session_state.recording) and st.session_state.stop_event.is_set():
        recorded_file_path = audio_filename
        st.audio(recorded_file_path)

    # ---------------------
    # File Upload Section
    # ---------------------
    st.subheader("Upload Audio File")
    uploaded_file = st.file_uploader(
        "Or upload an audio file (wav, mp3, m4a, ogg, flac)",
        type=["wav", "mp3", "m4a", "ogg", "flac"]
    )
    uploaded_file_path = None
    if uploaded_file is not None:
        with tempfile.NamedTemporaryFile(delete=False, suffix=uploaded_file.name) as tmp:
            tmp.write(uploaded_file.read())
            tmp.flush()
            uploaded_file_path = tmp.name
        st.audio(uploaded_file_path)

    # ---------------------
    # Transcription Section
    # ---------------------
    st.subheader("Transcription")
    audio_file_to_transcribe = uploaded_file_path if uploaded_file_path else recorded_file_path

    if audio_file_to_transcribe:
        if st.button("Transcribe Audio"):
            with st.spinner("Transcribing..."):
                # Load audio data and prepare for transcription.
                audio_data = whisper.load_audio(audio_file_to_transcribe)
                audio_data = whisper.pad_or_trim(audio_data)
                chosen_language = "sv" if language_mode == "Force Swedish" else None
                result = model.transcribe(audio_data, language=chosen_language, task=task)
            text = result.get("text", "")
            segments = result.get("segments", [])

            col1, col2 = st.columns(2)
            with col1:
                st.text_area("Full Transcript", text, height=300)
            with col2:
                srt_content = ""
                for i, segment in enumerate(segments, 1):
                    start_time = seconds_to_srt_time(segment["start"])
                    end_time = seconds_to_srt_time(segment["end"])
                    segment_text = segment["text"].strip()
                    st.markdown(f"**{start_time} - {end_time}**")
                    st.write(segment_text)
                    st.divider()
                    srt_content += f"{i}\n{start_time} --> {end_time}\n{segment_text}\n\n"
            st.download_button("Download SRT",
                               data=srt_content,
                               file_name="transcription.srt",
                               mime="text/plain")
            st.download_button("Download TXT",
                               data=text,
                               file_name="transcription.txt",
                               mime="text/plain")
            # Generate DOCX and provide as a download.
            word_file = generate_word_file(text)
            st.download_button("Download Word File",
                               data=word_file.getvalue(),
                               file_name="transcription.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    else:
        st.info("Please record or upload an audio file to transcribe.")

if __name__ == "__main__":
    main()
